from __future__ import annotations

import csv
import base64
import html
import json
import re
from collections import Counter
from dataclasses import asdict, dataclass
from pathlib import Path
from statistics import median
from typing import Iterable
from zipfile import ZipFile
import subprocess
import tempfile
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeout

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.style import WD_STYLE_TYPE
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}


@dataclass
class Issue:
    severity: str
    category: str
    location: str
    message: str
    actual: str = ""
    expected: str = ""
    suggestion: str = ""


def _safe(value, default="未设置"):
    return default if value is None else value


def _pt(value):
    return None if value is None else round(value.pt, 2)


def _cm(value):
    return None if value is None else round(value.cm, 2)


def _rgb(value):
    try:
        return str(value.rgb) if value and value.rgb else None
    except Exception:
        return None


def _first_effective_run(paragraph):
    return next((r for r in paragraph.runs if r.text.strip()), None)


def _font_signature(paragraph):
    run = _first_effective_run(paragraph)
    if not run:
        return {}
    east = run._element.rPr
    east_name = None
    if east is not None and east.rFonts is not None:
        east_name = east.rFonts.get(f"{{{W}}}eastAsia")
    font = run.font
    return {
        "latin": font.name,
        "east_asia": east_name,
        "size_pt": _pt(font.size),
        "bold": font.bold,
        "italic": font.italic,
        "color": _rgb(font.color),
    }


def _paragraph_signature(paragraph):
    fmt = paragraph.paragraph_format
    return {
        "style": paragraph.style.name if paragraph.style else "",
        "alignment": str(paragraph.alignment),
        "left_indent_cm": _cm(fmt.left_indent),
        "right_indent_cm": _cm(fmt.right_indent),
        "first_line_indent_cm": _cm(fmt.first_line_indent),
        "space_before_pt": _pt(fmt.space_before),
        "space_after_pt": _pt(fmt.space_after),
        "line_spacing": fmt.line_spacing if isinstance(fmt.line_spacing, (int, float)) else _pt(fmt.line_spacing),
        "font": _font_signature(paragraph),
    }


HEADING_PATTERNS = [
    (1, re.compile(r"^(第[一二三四五六七八九十百]+章|[一二三四五六七八九十]+、|绪论$|结论$|结语$|参考文献$|致谢$|附录)")),
    (2, re.compile(r"^(\d+\.\d+)(?!\.)|^[（(][一二三四五六七八九十]+[）)]")),
    (3, re.compile(r"^(\d+\.\d+\.\d+)|^\d+[.、](?!\d)")),
    (4, re.compile(r"^\d+\.\d+\.\d+\.\d+")),
]


def heading_level(paragraph):
    text = paragraph.text.strip()
    style = paragraph.style.name if paragraph.style else ""
    # Word 的大纲级别是真正的结构信息，优先于样式名称和编号猜测。
    ppr = paragraph._p.pPr
    if ppr is not None and ppr.outlineLvl is not None:
        try:
            return int(ppr.outlineLvl.val) + 1
        except (TypeError, ValueError):
            pass
    m = re.search(r"(?:Heading|标题)\s*(\d+)", style, re.I)
    if m:
        return int(m.group(1))
    for level, pattern in reversed(HEADING_PATTERNS):
        if pattern.search(text) and len(text) <= 80:
            # 单独的“1.”常见于问卷题号；含制表符、问号或以冒号结束时不按标题处理。
            if level == 3 and re.match(r"^\d+[.、](?!\d)", text):
                if "\t" in text or "?" in text or "？" in text or text.endswith(("：", ":")) or len(text) > 35:
                    continue
            return level
    return None


def _is_toc(paragraph):
    style = paragraph.style.name.lower() if paragraph.style else ""
    text = paragraph.text.strip()
    # 目录样式可能被用户改成正文样式；制表符/点引导线+页码是稳定的内容特征。
    return (style.startswith("toc") or style.startswith("目录") or
            bool(re.search(r"(?:\.{2,}|。{2,}|…{2,}|\t)\s*\d{1,4}$", text)))


def _normalize_heading(text):
    text = text.replace("［", "[").replace("］", "]").replace("．", ".")
    text = re.sub(r"(?:\t|\s+|\.{2,}|。{2,}|…{2,})\s*\d+$", "", text.strip())
    # WPS 目录常把编号、空格和中文标点处理得与正文不同，比较时只保留标题主体。
    text = re.sub(r"^(?:第\s*[一二三四五六七八九十百零]+\s*章|\d+(?:\.\d+)*|[一二三四五六七八九十百零]+)[、.．]\s*", "", text)
    return re.sub(r"\s+", "", text)


def _style_reference(doc: DocumentType):
    refs = {}
    samples = {1: [], 2: [], 3: [], 4: [], 0: []}
    for p in doc.paragraphs:
        if not p.text.strip() or _is_toc(p):
            continue
        # 模板格式样本只采信显式标题样式或 Word 大纲级别，避免正文编号被当成标题样本。
        ppr = p._p.pPr
        explicit = (ppr is not None and ppr.outlineLvl is not None) or bool(re.search(r"(?:Heading|标题)\s*\d+", p.style.name if p.style else "", re.I))
        level = heading_level(p) if explicit else None
        samples[level or 0].append(_paragraph_signature(p))
    for level, items in samples.items():
        if not items:
            continue
        serialized = [json.dumps(x, sort_keys=True, ensure_ascii=False) for x in items]
        refs[level] = json.loads(Counter(serialized).most_common(1)[0][0])
    return refs


def _comments(path: Path):
    results = []
    with ZipFile(path) as z:
        if "word/comments.xml" not in z.namelist():
            return results
        comments_root = etree.fromstring(z.read("word/comments.xml"))
        doc_root = etree.fromstring(z.read("word/document.xml"))
        by_id = {}
        for c in comments_root.xpath("//w:comment", namespaces=NS):
            cid = c.get(f"{{{W}}}id")
            by_id[cid] = "".join(c.xpath(".//w:t/text()", namespaces=NS)).strip()
        emitted = set()
        for start in doc_root.xpath("//w:commentRangeStart", namespaces=NS):
            cid = start.get(f"{{{W}}}id")
            p = start.getparent()
            while p is not None and p.tag != f"{{{W}}}p":
                p = p.getparent()
            anchor = "" if p is None else "".join(p.xpath(".//w:t/text()", namespaces=NS)).strip()
            results.append((cid, anchor[:100], by_id.get(cid, "")))
            emitted.add(cid)
        # 某些 WPS/Word 文档保留了批注正文但锚点已损坏，仍应报告，不能静默丢失。
        for cid, text in by_id.items():
            if cid not in emitted:
                results.append((cid, "批注锚点未找到", text))
    return results


def _comments_safe(path: Path, timeout_seconds=12):
    """批注 XML 偶尔异常膨胀；超时后返回明确状态，避免界面无限等待。"""
    pool = ThreadPoolExecutor(max_workers=1)
    future = pool.submit(_comments, path)
    try:
        return future.result(timeout=timeout_seconds)
    except FutureTimeout:
        future.cancel()
        return [("?", "批注解析超时", "文档批注结构过大或损坏，已跳过批注细节解析")]
    finally:
        pool.shutdown(wait=False, cancel_futures=True)


def _field_codes(path: Path):
    with ZipFile(path) as z:
        root = etree.fromstring(z.read("word/document.xml"))
        codes = root.xpath("//w:instrText/text()", namespaces=NS)
    joined = " ".join(codes).upper()
    return {"PAGE": joined.count("PAGE"), "TOC": joined.count("TOC"), "REF": joined.count(" REF ")}


def _footnote_info(path: Path):
    """读取脚注引用与脚注正文，避免只看正文角标。"""
    with ZipFile(path) as z:
        names = set(z.namelist())
        if "word/footnotes.xml" not in names:
            return {"references": 0, "notes": 0}
        root = etree.fromstring(z.read("word/footnotes.xml"))
        notes = [n for n in root.xpath("//w:footnote", namespaces=NS)
                 if n.get(f"{{{W}}}id") not in {"-1", "0"}]
        doc = etree.fromstring(z.read("word/document.xml"))
        refs = doc.xpath("//w:footnoteReference", namespaces=NS)
        return {"references": len(refs), "notes": len(notes)}


def _header_footer_xml_text(path: Path):
    with ZipFile(path) as z:
        result = []
        for name in z.namelist():
            if re.match(r"word/(header|footer)\d+\.xml$", name):
                root = etree.fromstring(z.read(name))
                text = "".join(root.xpath("//w:t/text()", namespaces=NS)).strip()
                if text:
                    result.append(text)
        return result


def _textbox_xml_text(path: Path):
    """读取 WPS 用文本框/形状承载的页眉页脚文字。"""
    with ZipFile(path) as z:
        result = []
        for name in z.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                try:
                    root = etree.fromstring(z.read(name))
                except Exception:
                    continue
                for box in root.xpath("//*[local-name()='txbxContent']"):
                    text = "".join(box.xpath(".//w:t/text()", namespaces=NS)).strip()
                    if text:
                        result.append(text)
        return result


def _header_footer_assets(path: Path):
    with ZipFile(path) as z:
        return [n for n in z.namelist() if re.match(r"word/(header|footer)\d+\.xml\.rels$", n)
                and b"/media/" in z.read(n)]


class ThesisAnalyzer:
    def __init__(self, template_path: str | Path, paper_path: str | Path):
        self.template_path = Path(template_path)
        self.paper_path = Path(paper_path)
        self._converted = []
        self.template_docx_path = self._docx_path(self.template_path)
        self.paper_docx_path = self._docx_path(self.paper_path)
        self.template = Document(self.template_docx_path)
        self.paper = Document(self.paper_docx_path)
        self.issues: list[Issue] = []
        self.template_styles = _style_reference(self.template)
        self.template_heading_texts = self._template_heading_texts()
        self.template_text_rules = any(re.search(r"(一级|二级|三级|四级|正文|页眉|页脚).{0,30}(字体|字号|加粗|居中|缩进|目录)", p.text, re.I) for p in self.template.paragraphs)
        self.template_comment_count = len(_comments_safe(self.template_docx_path))

    def _docx_path(self, path: Path):
        if path.suffix.lower() == ".docx":
            return path
        if path.suffix.lower() != ".doc":
            raise ValueError("仅支持 .doc 或 .docx 文件")
        # 旧版 .doc 通过本机已安装的 Microsoft Word 转为临时 docx，原文件不被修改。
        try:
            import win32com.client  # type: ignore
            fd, target = tempfile.mkstemp(suffix=".docx")
            Path(target).unlink(missing_ok=True)
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(str(path.resolve()), ReadOnly=True)
            doc.SaveAs2(str(target), FileFormat=16)
            doc.Close(False)
            word.Quit()
            self._converted.append(target)
            return Path(target)
        except Exception as exc:
            raise ValueError("无法读取 .doc：请在此电脑安装 Microsoft Word，或先另存为 .docx。") from exc

    def _template_heading_texts(self):
        result = {1: set(), 2: set(), 3: set(), 4: set()}
        for p in self.template.paragraphs:
            level = heading_level(p)
            text = _normalize_heading(p.text)
            if level in result and text and not _is_toc(p):
                result[level].add(text)
        return result

    def add(self, severity, category, location, message, actual="", expected="", suggestion=""):
        self.issues.append(Issue(severity, category, location, message, str(actual), str(expected), suggestion))

    def analyze(self):
        self.issues.clear()
        self._sections()
        self._headings()
        self._body_styles()
        self._toc()
        self._captions()
        self._citations()
        self._fields()
        self._headers_footers()
        self._footnotes()
        self._teacher_comments()
        self._status()
        order = {"严重": 0, "警告": 1, "提示": 2}
        self.issues.sort(key=lambda x: (order.get(x.severity, 9), x.category, x.location))
        return self.issues

    def _sections(self):
        t_sections, p_sections = self.template.sections, self.paper.sections
        if len(p_sections) != len(t_sections):
            self.add("提示", "页面设置", "全文", "论文分节数量与模板不同", len(p_sections), len(t_sections), "核对封面、目录、正文之间的分节符是否符合学校要求。")
        names = ["page_width", "page_height", "top_margin", "bottom_margin", "left_margin", "right_margin", "header_distance", "footer_distance"]
        labels = ["纸张宽度", "纸张高度", "上页边距", "下页边距", "左页边距", "右页边距", "页眉距离", "页脚距离"]
        for i, ps in enumerate(p_sections):
            ts = t_sections[min(i, len(t_sections) - 1)]
            for name, label in zip(names, labels):
                pv, tv = _cm(getattr(ps, name)), _cm(getattr(ts, name))
                if pv is not None and tv is not None and abs(pv - tv) > 0.08:
                    self.add("警告", "页面设置", f"第{i+1}节", f"{label}与模板不一致", f"{pv} cm", f"{tv} cm", "按模板调整本节页面设置。")

    def _headers_footers(self):
        def texts(doc):
            out = []
            for section in doc.sections:
                for part in (section.header, section.first_page_header, section.even_page_header,
                             section.footer, section.first_page_footer, section.even_page_footer):
                    out.extend(p.text.strip() for p in part.paragraphs if p.text.strip())
            return out
        t = list(dict.fromkeys(texts(self.template) + _header_footer_xml_text(self.template_docx_path) + _textbox_xml_text(self.template_docx_path)))
        p = list(dict.fromkeys(texts(self.paper) + _header_footer_xml_text(self.paper_docx_path) + _textbox_xml_text(self.paper_docx_path)))
        ta, pa = _header_footer_assets(self.template_docx_path), _header_footer_assets(self.paper_docx_path)
        if ta and not pa:
            self.add("警告", "页眉页脚", "页眉页脚", "模板页眉页脚包含图片或Logo，论文未识别到", "未识别", "应保留页眉页脚图片/Logo", "检查页眉图片是否被删除或转为浮动对象。")
        if t and not p:
            self.add("警告", "页眉页脚", "页眉页脚", "模板包含页眉或页脚文字，论文未识别到", "未识别", "应保留模板页眉页脚文字", "检查页眉页脚是否被删除或转为图片。")
        elif t:
            missing = [x for x in t if x not in p]
            if missing:
                self.add("警告", "页眉页脚", "页眉页脚", "论文缺少模板页眉页脚文字", "；".join(missing[:5]), "与模板文字一致", "检查不同节、首页和奇偶页页眉页脚内容。")

    def _footnotes(self):
        t, p = _footnote_info(self.template_docx_path), _footnote_info(self.paper_docx_path)
        # 模板和论文可以采用不同注释制度；只有论文实际出现脚注引用时才检查其完整性。
        if p["references"] and p["references"] > p["notes"]:
            self.add("警告", "脚注", "脚注", "脚注引用与脚注正文数量不一致", f"引用{p['references']}，正文{p['notes']}", "数量一致", "检查损坏或缺失的脚注定义。")

    def _headings(self):
        previous = 0
        counters = {}
        appendix_mode = False
        for i, p in enumerate(self.paper.paragraphs, 1):
            if not p.text.strip() or _is_toc(p):
                continue
            level = heading_level(p)
            if not level:
                continue
            if re.match(r"^(附录|附件|附表|附图)", p.text.strip()):
                appendix_mode = True
            if appendix_mode:
                # 附录内部按自身内容组织，不与正文标题树混合。
                previous = level
                continue
            loc = f"段落{i}：{p.text.strip()[:35]}"
            if previous and level > previous + 1 and level != 4:
                self.add("严重", "标题层级", loc, "标题层级发生跳级", f"上一级为{previous}级，本段为{level}级", "层级逐级递进", "检查是否缺少中间层级标题或标题级别设置错误。")
            previous = level
            expected = self.template_styles.get(level)
            actual = _paragraph_signature(p)
            ppr = p._p.pPr
            has_outline = ppr is not None and ppr.outlineLvl is not None
            if expected and not has_outline:
                mismatches = []
                for key in ["alignment", "first_line_indent_cm", "space_before_pt", "space_after_pt", "line_spacing"]:
                    av, ev = actual.get(key), expected.get(key)
                    if av is not None and ev is not None and (isinstance(av, (int,float)) and isinstance(ev, (int,float)) and abs(av-ev) > 0.08 or av != ev):
                        mismatches.append(key)
                af, ef = actual.get("font", {}), expected.get("font", {})
                for key in ["latin", "east_asia", "size_pt", "bold", "italic"]:
                    if af.get(key) is not None and ef.get(key) is not None and af.get(key) != ef.get(key):
                        mismatches.append("字体." + key)
                # 已明确设置 Word 大纲级别时，样式名不再作为判定依据；避免“样式不同但大纲正确”的误报。
                if mismatches and not has_outline:
                    self.add("警告", "标题格式", loc, "标题格式与模板同级标题的常用格式不一致", "、".join(mismatches), "模板同级标题格式", "核对字体、字号、加粗、缩进、对齐和段落间距。")
            style = p.style.name if p.style else ""
            has_outline = p._p.pPr is not None and p._p.pPr.outlineLvl is not None
            if style.lower() == "normal" and level <= 3 and not has_outline:
                self.add("提示", "标题格式", loc, "疑似标题使用了正文样式 Normal", style, f"模板{level}级标题样式", "应用对应标题样式，便于目录和交叉引用自动更新。")
            # 模板文字要求独立于模板中的段落位置/域：只比较同级标题文字集合。
            expected_texts = self.template_heading_texts.get(level, set())
            if expected_texts and _normalize_heading(p.text) not in expected_texts:
                self.add("提示", "标题文字", loc, "标题文字未按模板同级文字要求识别", _normalize_heading(p.text), "模板同级标题文字集合", "按学校模板的标题文字和编号规则核对，不要只依赖样式。")
            m = re.match(r"^(\d+(?:\.\d+){0,3})", p.text.strip())
            if m:
                parts = tuple(map(int, m.group(1).split(".")))
                parent = parts[:-1]
                key = (len(parts), parent)
                last = counters.get(key)
                if last is not None and parts[-1] != last + 1:
                    self.add("警告", "标题编号", loc, "同级标题编号可能不连续", parts[-1], last + 1, "核对是否漏号、重号或跨章节沿用了旧编号。")
                counters[key] = parts[-1]

    def _body_styles(self):
        """比较正文段落的常用样式属性，不把附录和目录混入正文。"""
        candidates = [p for p in self.template.paragraphs if p.text.strip() and not _is_toc(p) and not heading_level(p)]
        if not candidates:
            return
        expected = _paragraph_signature(candidates[0])
        checked = 0
        for i, p in enumerate(self.paper.paragraphs, 1):
            if checked >= 80 or not p.text.strip() or _is_toc(p) or heading_level(p):
                continue
            actual = _paragraph_signature(p)
            diffs = []
            for key in ("alignment", "first_line_indent_cm", "space_before_pt", "space_after_pt", "line_spacing"):
                av, ev = actual.get(key), expected.get(key)
                if av is not None and ev is not None and av != ev:
                    diffs.append(key)
            af, ef = actual.get("font", {}), expected.get("font", {})
            for key in ("size_pt", "bold", "italic", "east_asia"):
                if af.get(key) is not None and ef.get(key) is not None and af.get(key) != ef.get(key):
                    diffs.append("字体." + key)
            if diffs:
                self.add("警告", "正文段落样式", f"段落{i}", "正文段落样式与模板不一致", "、".join(diffs), "模板正文常用段落样式", "按模板正文段落的字体、缩进、对齐和间距调整。")
            checked += 1

    def _toc(self):
        toc = [_normalize_heading(p.text) for p in self.paper.paragraphs if _is_toc(p) and p.text.strip()]
        headings = []
        appendix_mode = False
        for p in self.paper.paragraphs:
            if not p.text.strip() or _is_toc(p):
                continue
            level = heading_level(p)
            if not level:
                continue
            if re.match(r"^(附录|附件|附表|附图)", p.text.strip()):
                appendix_mode = True
            if not appendix_mode:
                headings.append(_normalize_heading(p.text))
        if not toc:
            self.add("严重", "目录", "目录", "未识别到目录条目", "无", "目录应与正文标题对应", "插入或更新 Word 自动目录。")
            return
        toc_set, heading_set = set(toc), set(headings)
        for x in toc_set - heading_set:
            self.add("警告", "目录", "目录", "目录条目在正文标题中未找到", x, "正文存在对应标题", "更新目录，或检查正文标题文字是否已修改。")
        for x in heading_set - toc_set:
            if len(x) > 2:
                # 附录中的说明性小标题/正文标题通常不列入学校目录，除非目录中已有该条目。
                if x.startswith("附录") or getattr(self, "_is_appendix_heading", lambda _: False)(x):
                    continue
                self.add("提示", "目录", "正文", "正文标题未出现在目录中", x, "目录包含该级标题", "确认该标题是否应进入目录，并更新目录域。")

    def _is_appendix_heading(self, text):
        return bool(re.match(r"^(附录|附件|附表|附图)", text))

    def _captions(self):
        found = {"图": [], "表": []}
        for i, p in enumerate(self.paper.paragraphs, 1):
            text = p.text.strip()
            m = re.match(r"^(图|表)\s*(\d+)(?:[-－—.]?(\d+))?", text)
            # 题注通常是短独立段落；以“图X体现/表X显示”开头的正文说明不能当作题注。
            narrative = re.match(r"^(图|表)\s*\d+(?:[-－—.]?\d+)?\s*(体现|表明|显示|说明|反映|可见|可以)", text)
            if m and len(text) <= 70 and not narrative:
                found[m.group(1)].append((i, text, int(m.group(2)), int(m.group(3) or 0)))
        for kind, items in found.items():
            seen = set()
            for i, text, chapter, seq in items:
                key = (chapter, seq)
                if key in seen:
                    self.add("严重", "图表题注", f"段落{i}：{text[:35]}", f"{kind}题注编号重复", f"{chapter}-{seq}", "编号唯一", "重新编号并同步更新正文引用。")
                seen.add(key)
            by_chapter = {}
            for _, text, ch, seq in items:
                by_chapter.setdefault(ch, []).append((seq, text))
            for ch, values in by_chapter.items():
                # “图1”这类无分节编号不参与“1-1、1-2、2-1”连续性判断。
                seqs = sorted(x[0] for x in values if x[0] > 0)
                if seqs and seqs != list(range(1, max(seqs) + 1)):
                    self.add("警告", "图表题注", f"第{ch}章", f"{kind}题注编号不连续", seqs, list(range(1, max(seqs)+1)), "检查缺失或误编号的题注。")

    def _citations(self):
        # 参考文献表和目录中的编号不是正文引用；只在正文部分识别。
        paragraphs = list(self.paper.paragraphs)
        ref_start = next((i for i,p in enumerate(paragraphs) if p.text.strip() in {"参考文献", "参考文献表", "References"}), len(paragraphs))
        # 没有独立“参考文献”标题时，以连续的 [1]、[2] 条目定位文末参考文献区。
        def plain(s):
            return s.replace("［", "[").replace("］", "]").replace("﹝", "[").replace("﹞", "]")
        first_ref = next((i for i,p in enumerate(paragraphs) if re.match(r"^\s*\[\s*1\s*\]\s*", plain(p.text))), len(paragraphs))
        if first_ref < ref_start and first_ref > 0:
            ref_start = first_ref
        body = paragraphs[:ref_start]
        text = "\n".join(p.text for p in body if not _is_toc(p))
        refs = []
        for p in paragraphs[ref_start:]:
            m = re.match(r"^\s*\[\s*(\d+)\s*\]\s*", plain(p.text))
            if m:
                refs.append(int(m.group(1)))
        cites = []
        # 必须位于正文语境中：方括号内为数字且前后不是参考文献条目/目录页码。
        text = plain(text)
        # 兼容 WPS 全角括号、括号内空格及跨 run 后形成的常见变体。
        for m in re.finditer(r"(?<![A-Za-z0-9])\[\s*(\d+(?:\s*[-—,，]\s*\d+)*)\s*\](?!\d)", text):
            raw = m.group(1)
            for n in re.findall(r"\d+", raw):
                cites.append(int(n))
        if refs:
            refset, citeset = set(refs), set(cites)
            for n in sorted(citeset - refset):
                self.add("严重", "引用与参考文献", "正文", f"引用[{n}]在参考文献表中不存在", f"[{n}]", "文末存在对应文献", "补充文献或修正引用序号。")
            for n in sorted(refset - citeset):
                self.add("警告", "引用与参考文献", "参考文献", f"参考文献[{n}]未在正文中检出引用", f"[{n}]", "正文至少引用一次", "核对是否漏引或编号格式无法识别。")
            if refs != list(range(1, max(refs) + 1)):
                self.add("警告", "引用与参考文献", "参考文献", "参考文献编号不连续或顺序异常", refs, list(range(1, max(refs)+1)), "按正文首次引用顺序核对编号。")
        else:
            self.add("提示", "引用与参考文献", "参考文献", "未识别到以[1]开头的顺序编码参考文献", "无", "按学校要求判断", "若采用顺序编码制，请核对参考文献编号格式。")

    def _fields(self):
        t, p = _field_codes(self.template_docx_path), _field_codes(self.paper_docx_path)
        if t["PAGE"] and not p["PAGE"]:
            self.add("严重", "页码", "页眉页脚", "模板包含页码域，但论文未检出 PAGE 页码域", "未检出", "存在 PAGE 域", "检查页码是否被手工输入或误删。")
        if t["TOC"] and not p["TOC"]:
            self.add("警告", "目录", "目录", "模板使用自动目录域，但论文未检出 TOC 域", "未检出", "存在 TOC 域", "使用 Word 自动目录并在定稿前更新域。")

    def _teacher_comments(self):
        for cid, anchor, comment in _comments_safe(self.paper_docx_path):
            self.add("提示", "教师批注", f"批注{cid}：{anchor[:28]}", comment or "存在空批注", anchor, "按教师意见处理", "处理后在 Word 中回复或删除已解决批注。")

    def _status(self):
        severe = sum(1 for x in self.issues if x.severity == "严重")
        warnings = sum(1 for x in self.issues if x.severity == "警告")
        self.add("提示", "检查状态", "检查摘要", "模板已读取；教师批注已读取；格式检查结论：" + ("不合格" if severe or warnings else "合格"), f"严重{severe}，警告{warnings}", "严重和警告均为0时合格", "优先处理严重问题，再处理警告；提示项用于人工复核。")
        if self.template_text_rules and not self.template_styles:
            self.add("提示", "模板解析", "模板", "已读取文字说明型模板；标题层级按论文大纲级别/编号识别", "文字规范", "按文字说明执行", "WPS 中请为标题设置大纲级别，目录字体变化不会影响结构识别。")


def summary(issues: Iterable[Issue]):
    c = Counter(i.severity for i in issues)
    return {"严重": c["严重"], "警告": c["警告"], "提示": c["提示"], "总计": sum(c.values())}


def export_json(path, issues):
    s = summary(issues)
    Path(path).write_text(json.dumps({"report_type":"论文格式检查明细报告", "conclusion":"合格" if s["严重"] == 0 and s["警告"] == 0 else "不合格", "summary": s, "issues": [asdict(i) for i in issues]}, ensure_ascii=False, indent=2), encoding="utf-8")


def export_csv(path, issues):
    with open(path, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.writer(f)
        s = summary(issues)
        w.writerow(["报告类型", "结论", "严重程度", "类别", "位置", "问题", "当前情况", "模板要求", "修改建议"])
        for i in issues:
            w.writerow(["论文格式检查明细报告", "合格" if s["严重"] == 0 and s["警告"] == 0 else "不合格", i.severity, i.category, i.location, i.message, i.actual, i.expected, i.suggestion])


def export_html(path, issues, template_name="", paper_name="", logo_path=None):
    s = summary(issues)
    rows = "".join(
        "<tr>" + "".join(f"<td>{html.escape(str(v))}</td>" for v in [i.severity, i.category, i.location, i.message, i.actual, i.expected, i.suggestion]) + "</tr>"
        for i in issues
    )
    logo64 = ""
    if logo_path and Path(logo_path).exists():
        logo64 = base64.b64encode(Path(logo_path).read_bytes()).decode("ascii")
    watermark = f"body:after{{content:'';position:fixed;inset:20% 8%;background:url(data:image/png;base64,{logo64}) center/72% no-repeat;opacity:.035;z-index:-1}}" if logo64 else ""
    brand = f"<img alt='常青文创设计' src='data:image/png;base64,{logo64}' style='width:300px;max-height:85px;object-fit:contain'>" if logo64 else "<b>常青文创设计</b>"
    page = f"""<!doctype html><html lang='zh-CN'><head><meta charset='UTF-8'><meta http-equiv='Content-Type' content='text/html; charset=UTF-8'><title>常青文创论文格式检查报告</title>
<style>body{{font-family:'Microsoft YaHei','SimSun','Noto Sans CJK SC',sans-serif;margin:32px;color:#222}}{watermark}h1{{font-size:24px;color:#071952}}.cards{{display:flex;gap:12px;margin:18px 0}}.card{{padding:12px 20px;border-radius:8px;background:#f4f6f8}}table{{border-collapse:collapse;width:100%;font-size:13px}}th,td{{border:1px solid #d8dde3;padding:8px;vertical-align:top}}th{{background:#eef2f6;position:sticky;top:0}}td:first-child{{white-space:nowrap}}</style></head><body>
{brand}<h1>论文格式检查报告</h1><p>模板：{html.escape(template_name)}<br>论文：{html.escape(paper_name)}</p>
<div class='cards'><div class='card'>严重：{s['严重']}</div><div class='card'>警告：{s['警告']}</div><div class='card'>提示：{s['提示']}</div><div class='card'>总计：{s['总计']}</div></div>
<table><thead><tr><th>程度</th><th>类别</th><th>位置</th><th>问题</th><th>当前情况</th><th>模板要求</th><th>修改建议</th></tr></thead><tbody>{rows}</tbody></table></body></html>"""
    Path(path).write_text(page, encoding="utf-8")
